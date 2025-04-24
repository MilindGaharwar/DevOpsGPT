import yaml
from docx import Document
from PyPDF2 import PdfReader
import pandas as pd
from pathlib import Path
from typing import Dict, Any, Optional
from loguru import logger

class FileHandler:
    def __init__(self, workspace_path: Path):
        self.workspace = workspace_path

    def read_file(self, filepath: str) -> Dict[str, Any]:
        """Read various file types and return their content"""
        file_path = self.workspace / filepath
        if not file_path.exists():
            raise FileNotFoundError(f"File {filepath} not found")

        file_extension = file_path.suffix.lower()
        try:
            if file_extension == '.yaml' or file_extension == '.yml':
                return self._read_yaml(file_path)
            elif file_extension == '.docx':
                return self._read_docx(file_path)
            elif file_extension == '.pdf':
                return self._read_pdf(file_path)
            else:
                return self._read_text(file_path)
        except Exception as e:
            logger.error(f"Error reading file {filepath}: {str(e)}")
            raise

    def write_file(self, filepath: str, content: Any, file_type: Optional[str] = None) -> Dict[str, str]:
        """Write content to various file types"""
        file_path = self.workspace / filepath
        file_path.parent.mkdir(parents=True, exist_ok=True)

        try:
            if file_type == 'yaml' or filepath.endswith(('.yaml', '.yml')):
                return self._write_yaml(file_path, content)
            elif file_type == 'docx' or filepath.endswith('.docx'):
                return self._write_docx(file_path, content)
            else:
                return self._write_text(file_path, content)
        except Exception as e:
            logger.error(f"Error writing file {filepath}: {str(e)}")
            raise

    def analyze_yaml(self, content: str) -> Dict[str, Any]:
        """Analyze YAML content and provide suggestions"""
        try:
            yaml_data = yaml.safe_load(content)
            suggestions = []

            # Check for common YAML best practices
            if isinstance(yaml_data, dict):
                if not yaml_data.get('version'):
                    suggestions.append("Consider adding a 'version' field")
                if not yaml_data.get('description'):
                    suggestions.append("Consider adding a 'description' field")

            return {
                "valid": True,
                "structure": yaml_data,
                "suggestions": suggestions
            }
        except yaml.YAMLError as e:
            return {
                "valid": False,
                "error": str(e),
                "suggestions": ["Fix YAML syntax errors"]
            }

    def _read_yaml(self, file_path: Path) -> Dict[str, Any]:
        with open(file_path, 'r') as f:
            return {"content": yaml.safe_load(f), "type": "yaml"}

    def _read_docx(self, file_path: Path) -> Dict[str, Any]:
        doc = Document(file_path)
        content = [paragraph.text for paragraph in doc.paragraphs]
        return {"content": content, "type": "docx"}

    def _read_pdf(self, file_path: Path) -> Dict[str, Any]:
        reader = PdfReader(file_path)
        content = [page.extract_text() for page in reader.pages]
        return {"content": content, "type": "pdf"}

    def _read_text(self, file_path: Path) -> Dict[str, Any]:
        with open(file_path, 'r') as f:
            return {"content": f.read(), "type": "text"}

    def _write_yaml(self, file_path: Path, content: Any) -> Dict[str, str]:
        with open(file_path, 'w') as f:
            yaml.safe_dump(content, f, default_flow_style=False)
        return {"status": "success", "message": f"YAML file written successfully"}

    def _write_docx(self, file_path: Path, content: str) -> Dict[str, str]:
        doc = Document()
        doc.add_paragraph(content)
        doc.save(file_path)
        return {"status": "success", "message": "Word document written successfully"}

    def _write_text(self, file_path: Path, content: str) -> Dict[str, str]:
        with open(file_path, 'w') as f:
            f.write(content)
        return {"status": "success", "message": "Text file written successfully"}
